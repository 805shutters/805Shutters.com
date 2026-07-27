from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REGISTRY_PATH = ROOT / "src/lib/crm/vendor-orders/manufacturer-order-form-registry.json"
ORDER_TEMPLATE_ROOT = ROOT / "public/order-form-templates"
OUTPUT_ROOT = ROOT / "public/technical-measure-templates"
LOGO = ROOT / "public/brand/805-shutters-logo-paper.png"

# compact_reference_guide with a named "technician field form" geometry override.
# The visible table border starts at the first cell's text inset. The table
# width is reduced by the same amount so its right edge stays inside the
# section's 10,656-DXA content box in Word, LibreOffice, and Google Docs.
TABLE_INDENT_DXA = 85
PAGE_WIDTH_DXA = 10656 - TABLE_INDENT_DXA
ORANGE = "C56B2D"
BROWN = "74401F"
INK = "202020"
MUTED = "626262"
LIGHT = "F5F2EF"
PALE = "F8E9DE"
WHITE = "FFFFFF"
BORDER = "B8B8B8"
RED = "983B34"
GREEN = "39734B"

HEADER_FIELDS = {
    "customer_name",
    "crm_customer_id",
    "contract_number",
    "order_packet_id",
    "project_address",
    "customer_phone",
    "customer_email",
    "salesperson",
    "measure_status",
    "final_order_source",
    "line_item_number",
}
CORE_FIELDS = {
    "quantity",
    "room",
    "room_location",
    "source_opening_id",
    "side_mark_po",
    "ordered_width",
    "ordered_height",
    "width_a",
    "height_b",
    "mount_type",
}
CONTROL_FIELDS = {
    "compatibility_verification",
    "line_number",
    "portal_line_quantity",
}


def slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")


def title_slug(value: str) -> str:
    return "-".join(part.capitalize() for part in slug(value).split("-"))


def filename(entry: dict, extension: str = "docx") -> str:
    return (
        f"805-Shutters-{entry['manufacturer']}-"
        f"{title_slug(entry['product_name'])}-Technical-Measure-Form.{extension}"
    )


def field_label(key: str) -> str:
    exact = {
        "side_mark_po": "Side Mark / PO",
        "t_post_location": "T-Post Location",
        "top_down_bottom_up_option": "Top-Down / Bottom-Up",
        "width_a": "Width",
        "height_b": "Height",
    }
    return exact.get(key, key.replace("_", " ").title().replace(" Po", " PO"))


def set_font(run, size=7.6, bold=False, color=INK, italic=False, name="Arial"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(paragraph, before=0, after=0, line=1.0, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def border(cell, color=BORDER, size=5):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)
        node.set(qn("w:space"), "0")


def cell_margins(cell, top=55, start=85, bottom=55, end=85):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            value = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(value))
            tc_width.set(qn("w:type"), "dxa")
            border(cell)
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, widths):
    table = doc.add_table(rows=rows, cols=len(widths))
    set_table_geometry(table, widths)
    return table


def clear_cell(cell):
    cell.text = ""
    result = cell.paragraphs[0]
    paragraph(result)
    return result


def add_cell_text(cell, value, size=7.6, bold=False, color=INK, italic=False, align=None):
    para = clear_cell(cell)
    if align is not None:
        para.alignment = align
    run = para.add_run(value)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    return para


def add_band(doc, title, subtitle=""):
    table = add_table(doc, 1, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade(cell, BROWN)
    para = clear_cell(cell)
    run = para.add_run(title.upper())
    set_font(run, size=8.4, bold=True, color=WHITE)
    if subtitle:
        detail = para.add_run(f"   {subtitle}")
        set_font(detail, size=6.6, color=PALE)


def field_note(field: dict) -> str:
    options = field.get("options") or []
    if options:
        return "Options: " + " / ".join(str(item) for item in options)
    if field.get("input") == "boolean":
        return "[ ] Yes   [ ] No"
    if field.get("input") == "dimension":
        return "inches - nearest required fraction"
    if not field.get("required"):
        return "N/A when not applicable"
    return ""


def add_field_cell(cell, field):
    shade(cell, WHITE)
    para = clear_cell(cell)
    label_run = para.add_run(field["label"] + (" *" if field.get("required") else ""))
    set_font(label_run, size=6.3, bold=True, color=MUTED)
    value_para = cell.add_paragraph()
    paragraph(value_para, before=0, after=0)
    token = value_para.add_run("{{" + field["key"].upper() + "}}")
    set_font(token, size=7.4, bold=True, color=BROWN, name="Courier New")
    note = field_note(field)
    if note:
        note_para = cell.add_paragraph()
        paragraph(note_para)
        note_run = note_para.add_run(note)
        set_font(note_run, size=5.4, italic=True, color=MUTED)


def add_field_grid(doc, fields, columns=3):
    rows = max(1, (len(fields) + columns - 1) // columns)
    widths = [PAGE_WIDTH_DXA // columns] * columns
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    table = add_table(doc, rows, widths)
    for index in range(rows * columns):
        cell = table.cell(index // columns, index % columns)
        if index < len(fields):
            add_field_cell(cell, fields[index])
        else:
            shade(cell, LIGHT)
            add_cell_text(cell, "")


def schema_fields(entry):
    source = json.loads((ORDER_TEMPLATE_ROOT / entry["schema"]).read_text())
    if "ordered_fields" in source:
        keys = source["ordered_fields"]
        required = {"mount_type", "product_type", "shutter_category", "shutter_type"}
        properties = {}
    else:
        line = source.get("$defs", {}).get("line", {})
        keys = list(line.get("properties", {}).keys())
        required = set(line.get("required", []))
        properties = line.get("properties", {})
    fields = []
    for key in keys:
        if key in HEADER_FIELDS or key in CORE_FIELDS or key in CONTROL_FIELDS:
            continue
        property_definition = properties.get(key, {})
        enum = property_definition.get("enum")
        property_type = property_definition.get("type")
        if enum:
            input_type = "choice"
        elif property_type == "boolean":
            input_type = "boolean"
        elif any(marker in key for marker in ("width", "height", "length", "projection", "return", "location", "depth", "drop")):
            input_type = "dimension"
        elif key.endswith("_notes") or key.endswith("_instructions"):
            input_type = "long_text"
        else:
            input_type = "text"
        fields.append({
            "key": key,
            "label": field_label(key),
            "input": input_type,
            "required": key in required,
            "options": enum or [],
        })
    return fields


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.34)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.14)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(8)


def add_header(doc, entry):
    table = add_table(doc, 1, [2100, PAGE_WIDTH_DXA - 2100])
    left, right = table.rows[0].cells
    border(left, WHITE, 0)
    border(right, WHITE, 0)
    if LOGO.exists():
        para = clear_cell(left)
        image = para.add_run().add_picture(str(LOGO), width=Inches(1.05))
        image._inline.docPr.set("descr", "805 Shutters logo")
        image._inline.docPr.set("title", "805 Shutters")
    else:
        add_cell_text(left, "805 SHUTTERS", 10, True, ORANGE)
    para = clear_cell(right)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    kicker = para.add_run(f"{entry['manufacturer'].upper()} TECHNICAL MEASURE")
    set_font(kicker, size=7, bold=True, color=MUTED)
    title = right.add_paragraph()
    paragraph(title, align=WD_ALIGN_PARAGRAPH.RIGHT)
    title_run = title.add_run(entry["product_name"])
    set_font(title_run, size=17, bold=True, color=BROWN)
    subtitle = right.add_paragraph()
    paragraph(subtitle, align=WD_ALIGN_PARAGRAPH.RIGHT)
    subtitle_run = subtitle.add_run("ONE CONTRACT LINE - ONE PRODUCT-SPECIFIC MEASURE PAGE")
    set_font(subtitle_run, size=6.2, bold=True, color=ORANGE)


def add_route_strip(doc, entry):
    table = add_table(doc, 1, [2650, 4000, PAGE_WIDTH_DXA - 6650])
    values = [
        ("ROUTING KEY", entry["routing_key"]),
        ("ORDER DOCUMENT", Path(entry["template_docx"]).name),
        ("SOURCE", "Contract -> technician override"),
    ]
    for index, (label, value) in enumerate(values):
        cell = table.cell(0, index)
        shade(cell, PALE if index < 2 else LIGHT)
        para = clear_cell(cell)
        label_run = para.add_run(label)
        set_font(label_run, size=5.5, bold=True, color=MUTED)
        value_para = cell.add_paragraph()
        paragraph(value_para)
        value_run = value_para.add_run(value)
        set_font(value_run, size=6.3, bold=True, color=BROWN if index < 2 else INK, name="Courier New" if index == 0 else "Arial")


def build(entry):
    doc = Document()
    configure_document(doc)
    add_header(doc, entry)
    add_route_strip(doc, entry)
    add_band(doc, "Customer and contract source", "Prefilled from the signed contract")
    add_field_grid(doc, [
        {"key": "customer_name", "label": "Customer", "required": True},
        {"key": "contract_number", "label": "Contract Number", "required": True},
        {"key": "project_address", "label": "Project Address", "required": True},
    ])
    add_band(doc, "Opening measurements", "Technician values become authoritative when submitted")
    add_field_grid(doc, [
        {"key": "line_item_number", "label": "Contract Line", "required": True},
        {"key": "room_location", "label": "Room / Location", "required": True},
        {"key": "side_mark_po", "label": "Opening / Side Mark", "required": True},
        {"key": "quantity", "label": "Quantity", "required": True},
        {"key": "ordered_width", "label": "Measured Width", "required": True, "input": "dimension"},
        {"key": "ordered_height", "label": "Measured Height", "required": True, "input": "dimension"},
        {"key": "mount_type", "label": "Mount Type", "required": True},
        {"key": "measurement_basis", "label": "Measurement Basis", "required": True},
        {"key": "window_condition", "label": "Window / Opening Condition", "required": False},
    ])
    fields = schema_fields(entry)
    add_band(doc, "Product-specific ordering details", "Listed in the linked manufacturer ordering sequence")
    add_field_grid(doc, fields)
    add_band(doc, "Technician notes and release")
    add_field_grid(doc, [
        {"key": "technician_notes", "label": "Technician Notes", "required": False, "input": "long_text"},
        {"key": "exceptions", "label": "Exceptions / Conflicts", "required": False, "input": "long_text"},
        {"key": "measure_complete", "label": "Measure Complete", "required": True, "input": "boolean"},
        {"key": "technician_name", "label": "Technician", "required": True},
        {"key": "measured_at", "label": "Measured Date / Time", "required": True},
        {"key": "order_release_status", "label": "Order Release Status", "required": True},
    ])
    warning = add_table(doc, 1, [PAGE_WIDTH_DXA])
    cell = warning.cell(0, 0)
    shade(cell, "FCE8E6")
    para = clear_cell(cell)
    run = para.add_run("RELEASE GATE: ")
    set_font(run, size=6.2, bold=True, color=RED)
    detail = para.add_run("Do not order when a required field is blank, contradictory, unresolved, or incompatible with the linked product.")
    set_font(detail, size=6.2, bold=True, color=INK)
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        f"805 Shutters - {entry['routing_key']} - linked measure and ordering template - version 1"
    )
    set_font(run, size=5.5, color=MUTED)
    output_dir = OUTPUT_ROOT / entry["manufacturer"].lower()
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / filename(entry)
    doc.save(output)
    return output


def main():
    registry = json.loads(REGISTRY_PATH.read_text())
    entries = [
        entry
        for manufacturer_entries in registry["manufacturers"].values()
        for entry in manufacturer_entries
    ]
    outputs = [build(entry) for entry in entries]
    manifest = {
        "manifest_version": 1,
        "template_type": "product_specific_technical_measure",
        "line_item_rule": "one_contract_line_resolves_one_measure_template_and_one_order_template_by_routing_key",
        "templates": [
            {
                "routing_key": entry["routing_key"],
                "manufacturer": entry["manufacturer"],
                "product_key": entry["product_key"],
                "product_name": entry["product_name"],
                "technical_measure_docx": str(output.relative_to(ROOT / "public")),
                "technical_measure_pdf": str(output.with_suffix(".pdf").relative_to(ROOT / "public")),
                "order_template_docx": f"order-form-templates/{entry['template_docx']}",
                "order_schema": f"order-form-templates/{entry['schema']}",
            }
            for entry, output in zip(entries, outputs)
        ],
    }
    (OUTPUT_ROOT / "technical-measure-template-manifest.json").write_text(
        json.dumps(manifest, indent=2) + "\n"
    )
    print(json.dumps({
        "generated": len(outputs),
        "output_root": str(OUTPUT_ROOT),
        "manifest": str(OUTPUT_ROOT / "technical-measure-template-manifest.json"),
    }, indent=2))


if __name__ == "__main__":
    main()
